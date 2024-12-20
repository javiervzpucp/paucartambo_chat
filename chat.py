# -*- coding: utf-8 -*-
"""
Created on Thu Nov 14 11:08:21 2024
@author: jveraz
"""

import streamlit as st
from langchain_community.vectorstores import Vectara
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from dotenv import load_dotenv
from docx import Document  # Para crear documentos de Word

# Cargar las variables de entorno desde el archivo .env
load_dotenv()

# Configuración de credenciales
openai_api_key = st.secrets['openai']["OPENAI_API_KEY"]
vectara_customer_id = "2620549959"
vectara_corpus_id = 2
vectara_api_key = "zwt_nDJrR3X2jvq60t7xt0kmBzDOEWxIGt8ZJqloiQ"

# Validar que todas las variables se hayan cargado correctamente
if not openai_api_key:
    raise ValueError("Falta la API Key de OpenAI. Configúrala en el archivo .env")
if not vectara_customer_id or not vectara_corpus_id or not vectara_api_key:
    raise ValueError("Falta información de Vectara. Configúrala en el archivo .env")

# Configuración de Vectara como VectorStore
vectara = Vectara(
    vectara_customer_id=vectara_customer_id,
    vectara_corpus_id=vectara_corpus_id,
    vectara_api_key=vectara_api_key,
)

# Inicializar modelo LLM
llm = ChatOpenAI(
    model_name="gpt-4-turbo",
    temperature=0,
    openai_api_key=openai_api_key,
)

# Configuración de RAG con LangChain
retriever = vectara.as_retriever(search_kwargs={"k": 5})
qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    retriever=retriever,
    return_source_documents=True
)

# Interfaz de Streamlit
st.title("Prototipo de Chat sobre Devociones Marianas de Paucartambo")

# Mostrar imagen principal
st.image(
    "https://raw.githubusercontent.com/javiervzpucp/paucartambo/main/imagenes/1.png",
    caption="Virgen del Carmen de Paucartambo",
    use_container_width=True,
)

# Preguntas sugeridas
preguntas_sugeridas = [
    "¿Qué danzas se presentan en honor a la Mamacha Carmen?",
    "¿Cuál es el origen de las devociones marianas en Paucartambo?",
    "¿Qué papeles tienen los diferentes grupos de danza en la festividad?",
    "¿Cómo se celebra la festividad de la Virgen del Carmen?",
    "¿Cuál es el significado de las vestimentas en las danzas?",
]

# Botones para preguntas sugeridas
st.write("**Preguntas sugeridas:**")
if "query" not in st.session_state:
    st.session_state.query = ""

for pregunta in preguntas_sugeridas:
    if st.button(pregunta):
        st.session_state.query = pregunta  # Asignar directamente al estado de la sesión

# Entrada personalizada
query = st.text_input(
    "O pregunta algo relacionado con las Devociones Marianas de Paucartambo:",
    value=st.session_state.query,
)

# Botón para obtener respuestas
if st.button("Responder"):
    if query.strip():
        try:
            # Consultar el sistema RAG
            response = qa_chain({"query": query})
            answer = response["result"]

            # Filtrar documentos únicos por su metadato 'source'
            source_docs = response["source_documents"]
            unique_sources = {}
            for doc in source_docs:
                source = doc.metadata.get("source", "Fuente desconocida")
                if source not in unique_sources:
                    unique_sources[source] = doc.page_content[:300] + "..."

            st.write("**Respuesta generada:**")
            st.session_state["response_editable"] = st.text_area(
                "Editar la respuesta generada:",
                value=answer,
                height=200
            )

            st.write("**Documentos relacionados:**")
            for i, (source, content) in enumerate(unique_sources.items()):
                with st.expander(f"Fuente {i+1}: {source}"):
                    st.write(content)

            # Guardar la respuesta y la pregunta en la sesión
            st.session_state["last_query"] = query
        except Exception as e:
            st.error(f"Error: {e}")
    else:
        st.warning("Por favor, ingresa una pregunta válida.")

# Mostrar la última pregunta y respuesta generada si existen
if "response_editable" in st.session_state and "last_query" in st.session_state:
    st.write("**Última pregunta y respuesta generada:**")
    st.write(f"**Pregunta:** {st.session_state['last_query']}")
    st.write(f"**Respuesta editable:** {st.session_state['response_editable']}")

    # Exportar respuesta editable a un archivo Word
    def export_to_word(question, response):
        doc = Document()
        doc.add_heading("Respuesta Generada", level=1)
        doc.add_paragraph(f"Pregunta: {question}")
        doc.add_paragraph(f"Respuesta: {response}")
        file_path = "respuesta_generada.docx"
        doc.save(file_path)
        return file_path

    if st.button("Exportar respuesta a Word"):
        file_path = export_to_word(
            st.session_state["last_query"], st.session_state["response_editable"]
        )
        with open(file_path, "rb") as file:
            st.download_button(
                label="Descargar Respuesta",
                data=file,
                file_name="respuesta_generada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# Retroalimentación del usuario
st.write("**¿Esta respuesta fue útil?**")
col1, col2 = st.columns(2)

def save_to_vectara(query, response):
    """
    Guarda la consulta y respuesta útil en un documento de Vectara.
    """
    try:
        vectara.add_texts(
            texts=[
                f"Query: {query}\nResponse: {response}",
            ]
        )
        st.success("¡Respuesta marcada como útil y guardada en Vectara!")
    except Exception as e:
        st.error(f"Error al guardar la respuesta en Vectara: {e}")

with col1:
    if st.button("👍 Sí"):
        try:
            save_to_vectara(st.session_state["last_query"], st.session_state["response_editable"])
        except Exception as e:
            st.error(f"Error: {e}")

with col2:
    if st.button("👎 No"):
        st.warning("Gracias por tu retroalimentación. Trabajaremos para mejorar.")
